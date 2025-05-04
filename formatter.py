import os
import re
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn, OxmlElement

def add_page_number_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")

    run_page = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = "1"
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    run_page._r.extend([fld_char1, instr_text, fld_char2, fld_text, fld_char3])

    paragraph.add_run(" of ")

    run_total = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.text = "NUMPAGES"
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = "1"
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    run_total._r.extend([fld_char1, instr_text, fld_char2, fld_text, fld_char3])

def parse_docx_transcript(path):
    doc = Document(path)
    entries = []

    current_speaker = None
    speaker = None
    current_lines = []
    full_text = None

    skipped_first_date = False

    datetime_line_pattern = re.compile(
        r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4},?\s+\d{1,2}:\d{2}(?:\s*[APMapm]{2})?$"
    )

    for para in doc.paragraphs:
        line = para.text.strip()

        if not line or "stopped transcription" in line.lower():
            continue

        if not skipped_first_date and datetime_line_pattern.fullmatch(line):
            skipped_first_date = True
            continue

        match = re.match(r"^(.*)\s+(\d{1,2}:\d{2})\n?(.*)", line, re.DOTALL)

        if match:
            if speaker and current_lines:
                full_text = "\n".join(current_lines).strip()
                entries.append((speaker, timestamp, full_text))
                full_text = None
                current_lines = []

            speaker = match.group(1).strip()
            timestamp = match.group(2).strip()
            
            current_lines.append(match.group(3).strip())
        else:
            if speaker:
                current_lines.append(line)

    if speaker and current_lines:
        full_text = "\n".join(current_lines).strip()
        entries.append((speaker, timestamp, full_text))
    return entries

def parse_vtt_transcript(path):
    entries = []

    with open(path, "r", encoding="utf-8") as file:
        lines = [line.strip() for line in file if line.strip()]

    speaker = None
    current_speaker = None
    current_timestamp = None
    current_lines = []
    same_speaker = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # If this line is a timestamp
        if "-->" in line:
            timestamp = line.split("-->")[0].strip()
            mmss_match = re.search(r"(\d{2}:\d{2})", timestamp)
            timestamp = timestamp[-8:]  # format HH:MM:SS or MM:SS.SSS
            if mmss_match:
                if not same_speaker:
                    current_time = timestamp
                    same_speaker = True
            else:
                i += 1
                continue

            # Look ahead for <v Speaker> line
            i += 1
            if i < len(lines) and lines[i].startswith("<v "):
                speaker_line = lines[i]
                speaker_match = re.match(r"<v\s+([^>]+)>(.*)", speaker_line)
                if speaker_match:
                    speaker = speaker_match.group(1).strip()
                    text = speaker_match.group(2).strip().replace("</v>", "")

                    # If speaker changed, store previous block
                    if current_speaker != speaker and current_speaker and current_lines:
                        entries.append((current_speaker, current_timestamp, "\n".join(current_lines).strip()))
                        current_lines = []
                        current_timestamp = timestamp
                        current_time = timestamp
                    else:
                        current_timestamp = current_time

                    current_speaker = speaker
                    current_lines.append(text)

                    # Collect any additional lines until the next timestamp or speaker
                    i += 1
                    while i < len(lines) and not "-->" in lines[i] and not lines[i].startswith("<v "):
                        # current_lines.append(lines[i])
                        i += 1
                    continue
        i += 1

    # Save last block
    if current_speaker and current_lines:
        entries.append((current_speaker, current_timestamp, "\n".join(current_lines).strip()))

    return entries

def format_transcript(entries, coach_name=None, output_path=None):
    from docx import Document
    from docx.shared import Inches, RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import qn, OxmlElement
    import os

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Header with coach name
    header_para = section.header.paragraphs[0]
    if coach_name:
        header_run = header_para.add_run(coach_name)
        header_run.font.color.rgb = RGBColor(64, 64, 64)
        header_run.font.size = Pt(14)
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer with page numbers
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Page ")
    run_page = footer_para.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = "1"
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    run_page._r.extend([fld_char1, instr_text, fld_char2, fld_text, fld_char3])
    footer_para.add_run(" of ")
    run_total = footer_para.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.text = "NUMPAGES"
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = "1"
    fld_char3 = OxmlElement('w:fldCharType')
    fld_char3.set(qn('w:fldCharType'), 'end')
    run_total._r.extend([fld_char1, instr_text, fld_char2, fld_text, fld_char3])

    doc.add_heading("Coaching Session Transcript with Feedback", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add ICF legend
    legend_items = [
        ("SD", "Evidence of competency demonstration..."),
        ("LD", "Lack of evidence of demonstration..."),
        ("AMDOS", "Ask Me During Our Session"),
        ("SWMDOS", "Share With Me During Our Session"),
        ("CEQ", "Close Ended Question"),
        ("ECNN", "Expansive conversation not needed."),
        ("CD", "Cognitive Distortion"),
    ]
    for label, desc in legend_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph("")  # space
    doc.add_paragraph("")

    # Add table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].repeat_on_every_page = True

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Coaching Transcript'
    hdr_cells[1].text = "Mentor's Feedback"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, (speaker, timestamp, text) in enumerate(entries, start=1):
        row = table.add_row().cells
        para = row[0].paragraphs[0]
        para.add_run(f"{i} [").bold = False
        time_run = para.add_run(timestamp)
        time_run.font.color.rgb = RGBColor(105, 105, 105)
        para.add_run("] ")
        label = "Coach" if speaker == coach_name else "Client"
        bold = para.add_run(f"{label} {speaker}")
        bold.bold = True
        para.add_run(f" {text}")
        row[1].text = ""

    doc.add_paragraph("")
    doc.add_paragraph("Strengths:")
    doc.add_paragraph("")
    doc.add_paragraph("Progression Ideas:")

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

def main():
    file_path = input("Enter the path to the MS Teams .vtt or .docx transcript: ").strip()
    if not os.path.isfile(file_path):
        print("File not found.")
        return

    if file_path.lower().endswith(".docx"):
        entries = parse_docx_transcript(file_path)
    elif file_path.lower().endswith(".vtt"):
        entries = parse_vtt_transcript(file_path)
    else:
        print("Unsupported file type. Only .vtt and .docx are accepted.")
        return

    if not entries:
        print("No transcript entries parsed. Please check the file formatting.")
        return

    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_path = os.path.join(os.path.dirname(file_path), base_name + "_Formatted.docx")
    format_transcript(entries, output_path)

if __name__ == "__main__":
    main()
